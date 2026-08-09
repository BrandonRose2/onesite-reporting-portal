from __future__ import annotations

from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE_ROOT = Path("/home/ubuntu/manager-checklists")
OUTPUT_ROOT = Path("/home/ubuntu/manager-checklists-word")

NAVY = "122B4B"
TEAL = "0C7469"
LIGHT_TEAL = "EAF7F4"
LIGHT_GOLD = "FFF7E5"
GRAY = "F3F5F7"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=60, start=60, bottom=60, end=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="D6DBE1") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_text(paragraph, text: str, bold: bool = False, color: str | None = None, size: float | None = None) -> None:
    text = text.replace("\\|", "|").replace("`", "")
    # Preserve bracket checkboxes as editable visual markers, not form fields.
    fragments = re.split(r"(\*\*.*?\*\*)", text)
    for fragment in fragments:
        if not fragment:
            continue
        run = paragraph.add_run(fragment[2:-2] if fragment.startswith("**") and fragment.endswith("**") else fragment)
        run.bold = bold or (fragment.startswith("**") and fragment.endswith("**"))
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if size:
            run.font.size = Pt(size)


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip().replace("\\|", "|").replace("`", "") for cell in re.split(r"(?<!\\)\|", stripped)]


def is_divider_row(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if line.strip()]
    if len(rows) < 2:
        return
    header = rows[0]
    body = [row for row in rows[1:] if not is_divider_row("|" + "|".join(row) + "|")]
    columns = max(len(header), *(len(row) for row in body))
    table = document.add_table(rows=1, cols=columns)
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for index in range(columns):
        cell = header_cells[index]
        set_cell_fill(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_paragraph_text(paragraph, header[index] if index < len(header) else "", bold=True, color="FFFFFF", size=7.5)
    for row_number, row_values in enumerate(body):
        cells = table.add_row().cells
        for index in range(columns):
            cell = cells[index]
            if row_number % 2 == 1:
                set_cell_fill(cell, "FAFBFC")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            set_paragraph_text(paragraph, row_values[index] if index < len(row_values) else "", size=7.2)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05
    document.styles["Title"].font.name = "Aptos Display"
    document.styles["Title"].font.size = Pt(21)
    document.styles["Title"].font.color.rgb = RGBColor.from_string(NAVY)
    for style_name, size in (("Heading 1", 14), ("Heading 2", 11)):
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(TEAL)
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("BRANDON ROSE PORTFOLIO  |  CONFIDENTIAL MANAGER WORKING DOCUMENT")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated from supplied Availability and Delinquent & Prepaid source reports — validate current status before action.")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string("667085")


def convert_file(source_path: Path, target_path: Path) -> None:
    document = Document()
    style_document(document)
    lines = source_path.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if line.startswith("| ") and "|" in line[2:]:
            table_lines: list[str] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            set_paragraph_text(paragraph, line[2:], bold=True)
        elif line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            set_paragraph_text(paragraph, line[3:], bold=True)
        elif line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(5)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), LIGHT_TEAL)
            paragraph._p.get_or_add_pPr().append(shading)
            set_paragraph_text(paragraph, line[2:], color=NAVY)
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            set_paragraph_text(paragraph, line[2:])
        elif line.strip() == "---":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "D6DBE1")
            border.append(bottom)
            paragraph._p.get_or_add_pPr().append(border)
        elif line.strip():
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, line)
        index += 1
    document.save(target_path)


def main() -> None:
    if OUTPUT_ROOT.exists():
        shutil.rmtree(OUTPUT_ROOT)
    OUTPUT_ROOT.mkdir(parents=True)
    markdown_files = sorted(SOURCE_ROOT.glob("*_Manager_Checklist.md"))
    for source_path in markdown_files:
        convert_file(source_path, OUTPUT_ROOT / f"{source_path.stem}.docx")
    print(f"Generated {len(markdown_files)} editable Word checklists in {OUTPUT_ROOT}")


if __name__ == "__main__":
    main()
